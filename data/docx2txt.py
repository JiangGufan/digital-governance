# # scripts/10_ingest_convert.py
# import os, re, csv, sys
# from pathlib import Path
# from tqdm import tqdm
# from datetime import datetime
# from docx import Document

# RAW_DIR = Path("/project/data/raw_docx")
# DOC_TXT_DIR = Path("/project/data/txt_doc_level")
# META_DIR = Path("/project/data/meta")
# DOC_TXT_DIR.mkdir(parents=True, exist_ok=True)
# META_DIR.mkdir(parents=True, exist_ok=True)

# META_CSV = META_DIR / "doc_metadata.csv"

# def read_docx_text(docx_path: Path) -> str:
#     """尽量提取段落文本；去除空白、合并多余空行"""
#     doc = Document(str(docx_path))
#     paras = []
#     for p in doc.paragraphs:
#         t = p.text.strip()
#         # 跳过纯装饰性或模板性信息的弱规则（可按需扩充）
#         if not t:
#             continue
#         # 过滤常见打印页眉标识（可按需扩展）
#         if re.fullmatch(r"^第?\s*\d+\s*页$", t):
#             continue
#         paras.append(t)
#     # 合并为一个大文本，以双换行分段
#     text = "\n\n".join(paras)
#     # 统一全角空格等
#     text = re.sub(r"[ \t\u3000]+", " ", text)
#     # 合并 3+ 连续空行
#     text = re.sub(r"\n{3,}", "\n\n", text)
#     return text.strip()

# def infer_city_year(filename: str):
#     """从文件名弱解析城市/年份（若文件名含城市_年份_政府工作报告）"""
#     # 示例：内江_2024_政府工作报告.docx
#     city, year = None, None
#     m = re.search(r"([^\W_]+)[_\-](20\d{2})", filename)
#     if m:
#         city, year = m.group(1), m.group(2)
#     return city, year

# def main():
#     files = list(RAW_DIR.glob("*.docx"))
#     if not files:
#         print(f"No .docx found in {RAW_DIR}")
#         sys.exit(1)

#     with open(META_CSV, "w", newline="", encoding="utf-8") as f:
#         writer = csv.DictWriter(f, fieldnames=[
#             "doc_id","filename","city","year","bytes","mtime","txt_path"
#         ])
#         writer.writeheader()

#         for i, path in enumerate(tqdm(files, desc="DOCX->TXT")):
#             try:
#                 text = read_docx_text(path)
#                 if not text:
#                     print(f"[WARN] Empty after parse: {path.name}")
#                     continue

#                 doc_id = f"doc_{i:03d}"
#                 out_txt = DOC_TXT_DIR / f"{doc_id}.txt"
#                 out_txt.write_text(text, encoding="utf-8")

#                 st = path.stat()
#                 city, year = infer_city_year(path.name)
#                 writer.writerow({
#                     "doc_id": doc_id,
#                     "filename": path.name,
#                     "city": city or "",
#                     "year": year or "",
#                     "bytes": st.st_size,
#                     "mtime": datetime.fromtimestamp(st.st_mtime).isoformat(),
#                     "txt_path": str(out_txt)
#                 })
#             except Exception as e:
#                 print(f"[ERROR] {path.name}: {e}")

# if __name__ == "__main__":
#     main()


from pathlib import Path
from docx import Document
import re

# 输入和输出目录
INPUT_DIR = Path("docxV")             # 你的 150 个 docx 在这里
OUTPUT_DIR = Path("txt_doc_level")    # 输出 txt 的目录

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

def docx_to_text(docx_path: Path) -> str:
    """
    从 .docx 中抽取正文，简单去掉空行和多余空格。
    可以以后根据需要再加规则。
    """
    doc = Document(str(docx_path))
    paras = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        # 这里可以按需过滤一些页眉/页脚/模板话（先保留，后面再精细清洗）
        paras.append(text)

    # 用空行分段
    text = "\n\n".join(paras)
    # 合并多空格
    text = re.sub(r"[ \t\u3000]+", " ", text)
    # 合并 3 个以上连续空行
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

def main():
    docx_files = sorted(INPUT_DIR.glob("*.docx"))
    if not docx_files:
        print(f"No .docx files found in {INPUT_DIR.resolve()}")
        return

    print(f"Found {len(docx_files)} .docx files, converting...")

    for i, docx_path in enumerate(docx_files, 1):
        try:
            text = docx_to_text(docx_path)
            if not text:
                print(f"[WARN] Empty content after parsing: {docx_path.name}")
                continue

            # 输出文件名：和原文件同名，只是后缀改为 .txt
            out_path = OUTPUT_DIR / (docx_path.stem + ".txt")
            out_path.write_text(text, encoding="utf-8")
            print(f"[{i}/{len(docx_files)}] Saved: {out_path}")
        except Exception as e:
            print(f"[ERROR] {docx_path.name}: {e}")

if __name__ == "__main__":
    main()
